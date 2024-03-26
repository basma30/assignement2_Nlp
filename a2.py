import numpy as np
import pandas as pd
from collections import Counter
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import string
import nltk
import random

# Download NLTK resources
nltk.download('punkt')
nltk.download('wordnet')
nltk.download('stopwords')

# Generate documents
# documents = [
#     "Machine learning is a subfield of artificial intelligence.",
#     "The Python programming language is commonly used in machine learning projects.",
#     "Natural language processing helps computers understand human language.",
#     "Deep learning is a subset of machine learning that focuses on neural networks.",
#     "Data science combines statistics, programming, and domain knowledge.",
#     "Computer vision involves teaching computers to interpret and understand visual information."
# ]

from docx import Document
#searching a funcion to generate diffrent documents 
# Create a new Document
doc = Document()
doc.add_paragraph('Machine learning is a subfield of artificial intelligence.')
doc.save('d1.docx')

doc = Document()
doc.add_paragraph('The Python programming language is commonly used in machine learning projects.')
doc.save('d2.docx')

doc = Document()
doc.add_paragraph('Deep learning is a subset of machine learning that focuses on neural networks.')
doc.save('d3.docx')

documents =['d1.docx','d2.docx','d3.docx']

# Preprocessing functions
def clean_text(text):
    # Remove punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
    return text

def lemmatize_text(text):
    lemmatizer = WordNetLemmatizer()
    return ' '.join([lemmatizer.lemmatize(word) for word in word_tokenize(text)])

# Preprocess documents
cleaned_documents = [clean_text(doc.lower()) for doc in documents]
preprocessed_documents = [lemmatize_text(doc) for doc in cleaned_documents]

# Calculate TF for each document (from scratch)
tf_matrix = []
for document in preprocessed_documents:
    word_counts = Counter(word_tokenize(document))
    total_words = len(word_tokenize(document))
    tf = {word: count/total_words for word, count in word_counts.items()}
    tf_matrix.append(tf)

# Calculate IDF (from scratch)
idf = {}
total_docs = len(preprocessed_documents)
for document in preprocessed_documents:
    for word in set(word_tokenize(document)):
        idf[word] = idf.get(word, 0) + 1

for word in idf:
    idf[word] = np.log(total_docs / (idf[word] + 1))+1

# Calculate TF-IDF (from scratch)
tfidf_matrix = []
for tf in tf_matrix:
    tfidf = {word: tf[word] * idf[word] for word in tf}
    tfidf_matrix.append(tfidf)

# Normalize TF-IDF (from scratch)
for i, tfidf in enumerate(tfidf_matrix):
    norm_factor = np.sqrt(sum(tfidf[word] ** 2 for word in tfidf))
    for word in tfidf:
        tfidf[word] /= norm_factor
        if np.isnan(tfidf[word]=='NaN'):
             tfidf[word] = 0

# Get feature names (unique words)
feature_names = list(idf.keys())

# Create DataFrame for TF-IDF (from scratch)
df_tfidf_scratch = pd.DataFrame(tfidf_matrix, columns=feature_names)

# Display TF-IDF matrix (from scratch)
print("TF-IDF Matrix (from scratch):")
print(df_tfidf_scratch)

# TF-IDF using scikit-learn
# TF-IDF vectorizer
tfidf_vectorizer = TfidfVectorizer()

# Fit TF-IDF vectorizer and transform documents
tfidf_matrix_sklearn = tfidf_vectorizer.fit_transform(preprocessed_documents)

# Get feature names (unique words)
feature_names_sklearn = tfidf_vectorizer.get_feature_names_out()

# Create DataFrame for TF-IDF (using scikit-learn)
df_tfidf_sklearn = pd.DataFrame(tfidf_matrix_sklearn.toarray(), columns=feature_names_sklearn)

# Display TF-IDF matrix (using scikit-learn)
print("\nTF-IDF Matrix (using scikit-learn):")
print(df_tfidf_sklearn)


















